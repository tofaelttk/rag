#!/usr/bin/env python3
"""
rag_cli.py — Terminal RAG: LlamaIndex + Pinecone + Ollama/Llama3 + HuggingFace BGE
Usage:
    python rag_cli.py                        # interactive mode (asks for files)
    python rag_cli.py path/to/file.pdf       # ingest one file then chat
    python rag_cli.py doc1.pdf doc2.txt      # ingest multiple files then chat
    python rag_cli.py --reuse                # skip ingestion, reuse existing Pinecone index
"""

import os
import sys
import time
import warnings
import textwrap
from pathlib import Path

# ── silence noisy logs before heavy imports ───────────────────────────────────
warnings.filterwarnings("ignore")
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["TRANSFORMERS_VERBOSITY"] = "error"
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

# ── Apple Silicon MPS fallback ────────────────────────────────────────────────
if os.getenv("FORCE_CPU_EMBEDDINGS", "0").lower() in ("1", "true", "yes"):
    os.environ["PYTORCH_ENABLE_MPS_FALLBACK"] = "1"
    try:
        import torch
        if hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
            torch.backends.mps.enabled = False
    except ImportError:
        pass

from dotenv import load_dotenv
load_dotenv()

# ── ANSI colours ──────────────────────────────────────────────────────────────
RESET  = "\033[0m"
BOLD   = "\033[1m"
DIM    = "\033[2m"
CYAN   = "\033[36m"
GREEN  = "\033[32m"
YELLOW = "\033[33m"
RED    = "\033[31m"
BLUE   = "\033[34m"
PURPLE = "\033[35m"
WHITE  = "\033[97m"

def c(color, text):   return f"{color}{text}{RESET}"
def tag(label, color): return c(color + BOLD, f"[{label}]")


# ── config ────────────────────────────────────────────────────────────────────
PINECONE_INDEX = os.getenv("PINECONE_INDEX_NAME", "rag-demo")
EMBED_MODEL    = os.getenv("EMBEDDING_MODEL",     "BAAI/bge-small-en-v1.5")
EMBED_DIM      = 384
CHUNK_SIZE     = int(os.getenv("CHUNK_SIZE",    "512"))
CHUNK_OVERLAP  = int(os.getenv("CHUNK_OVERLAP", "50"))
OLLAMA_MODEL   = os.getenv("OLLAMA_MODEL",  "llama3")
OLLAMA_URL     = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
TOP_K          = int(os.getenv("TOP_K", "5"))


# ── helpers ───────────────────────────────────────────────────────────────────
def hr(char="─", width=60):
    return c(DIM, char * width)

def log(label, msg, color=CYAN):
    ts = time.strftime("%H:%M:%S")
    print(f"  {c(DIM, ts)}  {tag(label, color)}  {msg}")

def log_ok(label, msg):    log(label, c(GREEN, msg),  GREEN)
def log_warn(label, msg):  log(label, c(YELLOW, msg), YELLOW)
def log_err(label, msg):   log(label, c(RED, msg),    RED)
def log_dim(label, msg):   log(label, c(DIM, msg),    DIM)


def print_banner():
    print()
    print(c(BOLD + WHITE, "  ╔══════════════════════════════════════════════════════╗"))
    print(c(BOLD + WHITE, "  ║") + c(BOLD + CYAN,  "      RAG CLI  —  Terminal Document Intelligence      ") + c(BOLD + WHITE, "║"))
    print(c(BOLD + WHITE, "  ╠══════════════════════════════════════════════════════╣"))
    print(c(BOLD + WHITE, "  ║") + c(DIM,           "  Stack : LlamaIndex · Pinecone · Ollama · BGE Embed  ") + c(BOLD + WHITE, "║"))
    print(c(BOLD + WHITE, "  ╚══════════════════════════════════════════════════════╝"))
    print()


def print_config():
    print(hr())
    log_dim("CONFIG", f"Pinecone index  : {PINECONE_INDEX}")
    log_dim("CONFIG", f"Embed model     : {EMBED_MODEL}  dim={EMBED_DIM}")
    log_dim("CONFIG", f"LLM             : {OLLAMA_MODEL}  @  {OLLAMA_URL}")
    log_dim("CONFIG", f"Chunk size      : {CHUNK_SIZE}  overlap={CHUNK_OVERLAP}")
    log_dim("CONFIG", f"Top-K           : {TOP_K}")
    print(hr())
    print()


def check_dependencies():
    """Verify Pinecone key, Ollama is running, etc."""
    ok = True

    # Pinecone key
    if not os.getenv("PINECONE_API_KEY"):
        log_err("CHECK", "PINECONE_API_KEY not set — add it to .env")
        ok = False
    else:
        log_ok("CHECK", "Pinecone API key  ✓")

    # Ollama
    try:
        import httpx
        r = httpx.get(f"{OLLAMA_URL}/api/tags", timeout=4)
        if r.status_code == 200:
            models = [m.get("name","") for m in r.json().get("models", [])]
            model_found = any(m.startswith(OLLAMA_MODEL) for m in models)
            if model_found:
                log_ok("CHECK", f"Ollama running  ·  {OLLAMA_MODEL} found  ✓")
            else:
                log_warn("CHECK", f"Ollama running but {OLLAMA_MODEL} not found")
                log_warn("CHECK", f"  Run:  ollama pull {OLLAMA_MODEL}")
                log_warn("CHECK", f"  Available: {', '.join(models) or 'none'}")
                ok = False
        else:
            raise Exception("bad status")
    except Exception:
        log_err("CHECK", f"Ollama not reachable at {OLLAMA_URL}")
        log_err("CHECK", "  Run:  ollama serve")
        ok = False

    return ok


# ── lazy singletons ───────────────────────────────────────────────────────────
_pinecone_client = None
_embed_model     = None
_llm             = None

def get_pinecone():
    global _pinecone_client
    if _pinecone_client is None:
        from pinecone import Pinecone
        _pinecone_client = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
    return _pinecone_client

def get_embed():
    global _embed_model
    if _embed_model is None:
        log("EMBED", f"Loading {EMBED_MODEL}…", PURPLE)
        from llama_index.embeddings.huggingface import HuggingFaceEmbedding
        _embed_model = HuggingFaceEmbedding(model_name=EMBED_MODEL, trust_remote_code=True)
        log_ok("EMBED", f"Model loaded  ✓")
    return _embed_model

def get_llm():
    global _llm
    if _llm is None:
        from llama_index.llms.ollama import Ollama
        _llm = Ollama(model=OLLAMA_MODEL, base_url=OLLAMA_URL, request_timeout=180.0, temperature=0.1)
    return _llm

def configure_settings():
    from llama_index.core import Settings
    Settings.embed_model   = get_embed()
    Settings.llm           = get_llm()
    Settings.chunk_size    = CHUNK_SIZE
    Settings.chunk_overlap = CHUNK_OVERLAP

def ensure_pinecone_index():
    from pinecone import ServerlessSpec
    pc = get_pinecone()
    existing = [i.name for i in pc.list_indexes()]
    if PINECONE_INDEX not in existing:
        log("PINECONE", f"Creating index '{PINECONE_INDEX}'…", GREEN)
        pc.create_index(
            name=PINECONE_INDEX, dimension=EMBED_DIM, metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1"),
        )
        log_ok("PINECONE", f"Index '{PINECONE_INDEX}' created  ✓")
    else:
        log_ok("PINECONE", f"Index '{PINECONE_INDEX}' exists  ✓")
    return pc.Index(PINECONE_INDEX)

def get_vector_store():
    from llama_index.vector_stores.pinecone import PineconeVectorStore
    return PineconeVectorStore(pinecone_index=ensure_pinecone_index())


# ── text extraction ───────────────────────────────────────────────────────────
def extract_text(path: Path):
    """Returns list of (text, metadata) tuples."""
    ext = path.suffix.lower()
    results = []

    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        results.append((text, {"source": path.name, "page": i + 1}))
        except ImportError:
            from llama_index.core import SimpleDirectoryReader
            docs = SimpleDirectoryReader(input_files=[str(path)], filename_as_id=True).load_data()
            results = [(d.text, {"source": path.name}) for d in docs if d.text.strip()]

    elif ext in (".docx", ".doc"):
        import docx as _docx
        doc  = _docx.Document(path)
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if text.strip():
            results.append((text, {"source": path.name}))

    elif ext in (".txt", ".md", ".rst"):
        text = path.read_text(encoding="utf-8", errors="replace")
        if text.strip():
            results.append((text, {"source": path.name}))

    else:
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
            if text.strip():
                results.append((text, {"source": path.name}))
        except Exception:
            raise ValueError(f"Cannot read file: {path.name}")

    return results


# ── ingest ────────────────────────────────────────────────────────────────────
def ingest(file_paths: list[Path]) -> bool:
    from llama_index.core import VectorStoreIndex, StorageContext
    from llama_index.core.node_parser import SentenceSplitter
    from llama_index.core.schema import Document

    print()
    print(hr("═"))
    log("INGEST", f"Starting ingestion of {len(file_paths)} file(s)", CYAN)
    print(hr("═"))
    print()

    configure_settings()
    all_nodes = []
    embed = get_embed()

    for fp in file_paths:
        if not fp.exists():
            log_err("FILE", f"Not found: {fp}")
            continue

        size_kb = fp.stat().st_size / 1024
        print()
        log("FILE", c(YELLOW, f"── {fp.name} ──────────────────────────────"), YELLOW)
        log_dim("FILE", f"Path : {fp.resolve()}")
        log_dim("FILE", f"Size : {size_kb:.1f} KB")

        # Extract
        log("PARSE", "Extracting text…", PURPLE)
        t0 = time.time()
        try:
            segments = extract_text(fp)
        except Exception as e:
            log_err("PARSE", str(e))
            continue

        word_count = sum(len(t.split()) for t, _ in segments)
        log_ok("PARSE", f"{len(segments)} segment(s)  ·  {word_count:,} words  ({time.time()-t0:.2f}s)")

        if not segments:
            log_warn("PARSE", "No text found — skipping")
            continue

        # Build Documents
        docs = [Document(text=t, metadata=m, doc_id=f"{m['source']}::s{i}")
                for i, (t, m) in enumerate(segments)]

        # Chunk
        log("CHUNK", f"Sentence-splitting  (size={CHUNK_SIZE}  overlap={CHUNK_OVERLAP})…", GREEN)
        splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        nodes    = splitter.get_nodes_from_documents(docs, show_progress=False)

        # Filter empty
        valid, skipped = [], 0
        for n in nodes:
            text = n.get_content() if hasattr(n, "get_content") else str(n)
            if text and text.strip() and len(text.strip()) > 15:
                valid.append(n)
            else:
                skipped += 1
        log_ok("CHUNK", f"{len(valid)} valid chunks  ({skipped} empty skipped)")

        if not valid:
            log_warn("CHUNK", "No valid chunks — skipping file")
            continue

        # Validate embeddings
        log("EMBED", f"Validating {len(valid)} embeddings…", PURPLE)
        t0     = time.time()
        good   = []
        bad    = 0
        total  = len(valid)

        for i, node in enumerate(valid):
            text = node.get_content() if hasattr(node, "get_content") else str(node)
            try:
                emb = embed.get_text_embedding(text)
                if emb and any(abs(float(v)) > 1e-10 for v in emb):
                    good.append(node)
                else:
                    bad += 1
            except Exception:
                bad += 1

            # progress bar
            pct  = int(((i + 1) / total) * 100)
            done = int(pct / 2)
            bar  = c(BLUE, "█" * done) + c(DIM, "░" * (50 - done))
            print(f"\r    {bar}  {pct:3d}%  ({i+1}/{total})", end="", flush=True)

        print()  # newline after progress bar
        elapsed = time.time() - t0
        log_ok("EMBED", f"{len(good)} valid  ·  {bad} zero/invalid skipped  ({elapsed:.2f}s)")

        if not good:
            log_err("EMBED", "All embeddings were zero — skipping file")
            continue

        all_nodes.extend(good)

    if not all_nodes:
        log_err("INGEST", "No valid nodes produced across all files")
        return False

    # Upsert to Pinecone
    print()
    log("PINECONE", f"Connecting to index '{PINECONE_INDEX}'…", GREEN)
    t0 = time.time()
    try:
        vs  = get_vector_store()
        ctx = StorageContext.from_defaults(vector_store=vs)
        log("PINECONE", f"Upserting {len(all_nodes)} vectors…", GREEN)
        VectorStoreIndex(nodes=all_nodes, storage_context=ctx, show_progress=False)
        elapsed = time.time() - t0
        log_ok("PINECONE", f"Upsert complete  ({elapsed:.2f}s)  ✓")
    except Exception as e:
        log_err("PINECONE", f"Upsert failed: {e}")
        return False

    print()
    print(hr("═"))
    log_ok("DONE", f"Indexed {len(all_nodes)} vectors across {len(file_paths)} file(s)  ✓")
    print(hr("═"))
    print()
    return True


# ── query loop ────────────────────────────────────────────────────────────────
def query_loop():
    from llama_index.core import VectorStoreIndex

    configure_settings()

    log("INDEX", f"Loading index from Pinecone '{PINECONE_INDEX}'…", CYAN)
    try:
        vs    = get_vector_store()
        index = VectorStoreIndex.from_vector_store(vector_store=vs)
        log_ok("INDEX", "Index ready  ✓")
    except Exception as e:
        log_err("INDEX", f"Could not load index: {e}")
        sys.exit(1)

    print()
    print(hr("═"))
    print(c(BOLD + CYAN, "  INTERACTIVE MODE") + c(DIM, "  —  type your question, Enter to send"))
    print(c(DIM,         "  Commands: /exit  /quit  /q  to quit"))
    print(c(DIM,         "            /sources      to toggle source display"))
    print(c(DIM,         "            /top N        to change Top-K (e.g. /top 8)"))
    print(hr("═"))
    print()

    show_sources = True
    top_k        = TOP_K
    history      = []

    while True:
        try:
            print(c(BOLD + GREEN, "  You ❯ "), end="", flush=True)
            question = input().strip()
        except (KeyboardInterrupt, EOFError):
            print()
            print(c(YELLOW, "\n  Goodbye!\n"))
            break

        if not question:
            continue

        # commands
        if question.lower() in ("/exit", "/quit", "/q"):
            print(c(YELLOW, "\n  Goodbye!\n"))
            break
        if question.lower() == "/sources":
            show_sources = not show_sources
            print(c(DIM, f"  Source display: {'ON' if show_sources else 'OFF'}\n"))
            continue
        if question.lower().startswith("/top "):
            try:
                top_k = int(question.split()[1])
                print(c(DIM, f"  Top-K set to {top_k}\n"))
            except (IndexError, ValueError):
                print(c(DIM, "  Usage: /top 8\n"))
            continue
        if question.lower() == "/help":
            print(c(DIM, "  /exit     quit"))
            print(c(DIM, "  /sources  toggle source chunks"))
            print(c(DIM, "  /top N    set retrieval Top-K"))
            print()
            continue

        # RAG
        print()
        t0 = time.time()

        try:
            query_engine = index.as_query_engine(
                similarity_top_k=top_k,
                streaming=True,
            )

            log("QUERY", c(DIM, question[:80] + ("…" if len(question) > 80 else "")), CYAN)
            log_dim("RETRIEVE", f"Fetching top-{top_k} chunks from Pinecone…")

            response = query_engine.query(question)

            print()
            print(c(BOLD + GREEN, "  Assistant ❯"))
            print()

            full_text = ""
            # Stream with word-wrap
            for chunk in response.response_gen:
                print(chunk, end="", flush=True)
                full_text += chunk

            elapsed = time.time() - t0
            print()
            print()

            # Sources
            if show_sources and hasattr(response, "source_nodes") and response.source_nodes:
                print(hr("┄"))
                print(c(DIM + BOLD, "  Retrieved chunks:"))
                for i, sn in enumerate(response.source_nodes):
                    meta  = getattr(sn.node, "metadata", {})
                    src   = meta.get("source", meta.get("file_name", "unknown"))
                    page  = meta.get("page", "")
                    score = f"{sn.score:.4f}" if sn.score else "—"
                    page_str = f"  p.{page}" if page else ""
                    excerpt = sn.node.text[:120].replace("\n", " ")
                    print(c(DIM, f"    [{i+1}]  {src}{page_str}  score={score}"))
                    print(c(DIM, f"         "{excerpt}…""))
                print(hr("┄"))

            print(c(DIM, f"  ⏱  {elapsed:.2f}s  ·  top-{top_k}  ·  {OLLAMA_MODEL}  ·  pinecone:{PINECONE_INDEX}"))
            print()

        except Exception as e:
            err = str(e).lower()
            print()
            log_err("ERROR", str(e))
            if "timeout" in err or "timed out" in err:
                log_warn("TIP", "Ollama timed out — model may still be loading, try again")
            elif "connection" in err:
                log_warn("TIP", f"Ollama not reachable — run: ollama serve")
            print()


# ── file picker (interactive) ─────────────────────────────────────────────────
def interactive_file_picker() -> list[Path]:
    print(c(BOLD + CYAN, "  ┌─ Document Ingestion ───────────────────────────────────┐"))
    print(c(BOLD + CYAN, "  │") + c(DIM, "  Enter file paths one by one. Leave blank when done.  ") + c(BOLD + CYAN, "│"))
    print(c(BOLD + CYAN, "  │") + c(DIM, "  Supports: PDF · TXT · DOCX · MD · RST               ") + c(BOLD + CYAN, "│"))
    print(c(BOLD + CYAN, "  └────────────────────────────────────────────────────────┘"))
    print()

    paths = []
    idx   = 1
    while True:
        try:
            print(c(YELLOW, f"  File {idx} path") + c(DIM, " (or Enter to finish): "), end="", flush=True)
            raw = input().strip()
        except (KeyboardInterrupt, EOFError):
            break

        if not raw:
            if not paths:
                print(c(DIM, "  No files entered. Exiting.\n"))
                sys.exit(0)
            break

        # expand ~ and resolve
        p = Path(raw).expanduser().resolve()
        if not p.exists():
            print(c(RED, f"  ✗ Not found: {p}"))
            continue
        if p in paths:
            print(c(DIM, f"  Already added: {p.name}"))
            continue

        paths.append(p)
        print(c(GREEN, f"  ✓ Added: {p.name}  ({p.stat().st_size/1024:.1f} KB)"))
        idx += 1

    return paths


# ── main ──────────────────────────────────────────────────────────────────────
def main():
    print_banner()
    print_config()

    # Parse args
    args     = sys.argv[1:]
    reuse    = "--reuse" in args
    file_args = [Path(a).expanduser().resolve() for a in args if not a.startswith("--")]

    if reuse:
        log("MODE", "Reusing existing Pinecone index — skipping ingestion", CYAN)
        print()
        if not check_dependencies():
            sys.exit(1)
        query_loop()
        return

    if not check_dependencies():
        sys.exit(1)

    # Collect files
    if file_args:
        missing = [str(p) for p in file_args if not p.exists()]
        if missing:
            for m in missing:
                log_err("FILE", f"Not found: {m}")
            sys.exit(1)
        file_paths = file_args
        print()
        print(c(BOLD, "  Files to ingest:"))
        for p in file_paths:
            print(c(DIM, f"    • {p.name}  ({p.stat().st_size/1024:.1f} KB)"))
        print()
    else:
        file_paths = interactive_file_picker()
        print()

    # Ingest
    ok = ingest(file_paths)
    if not ok:
        log_err("MAIN", "Ingestion failed — nothing to query")
        sys.exit(1)

    # Query
    query_loop()


if __name__ == "__main__":
    main()
